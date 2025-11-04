#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文件处理模块
负责将各种格式的输入文件转换为统一格式的图片
"""

import os
from pathlib import Path
from typing import List, Tuple
from PIL import Image
import numpy as np

try:
    from pdf2image import convert_from_path
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("警告: pdf2image 未安装，无法处理PDF文件")

try:
    from paddleocr import PaddleOCR
    PADDLEOCR_AVAILABLE = True
except ImportError:
    PADDLEOCR_AVAILABLE = False
    print("警告: PaddleOCR 未安装")

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("提示: python-docx 未安装，安装后可处理 DOCX 文件")
    print("     运行: pip install python-docx")

try:
    import win32com.client
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False
    # DOC格式需要COM支持，只在需要时提示


class FileHandler:
    """文件处理类"""
    
    def __init__(self, config=None):
        """
        初始化文件处理器
        
        Args:
            config: 配置对象
        """
        self.config = config
        self.supported_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.bmp', '.gif', '.docx']
        if DOCX_SUPPORT:
            print(f"[文件处理] 支持 DOCX 格式")
        if DOC_SUPPORT:
            self.supported_extensions.append('.doc')
            print(f"[文件处理] 支持 DOC 格式")
    
    def convert_to_images(self, input_path: str) -> List[Image.Image]:
        """
        将输入文件转换为图片列表
        
        Args:
            input_path: 输入文件路径
            
        Returns:
            图片对象列表
        """
        file_path = Path(input_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {input_path}")
        
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            return self._convert_pdf(input_path)
        elif file_ext == '.docx':
            return self._convert_docx(input_path)
        elif file_ext == '.doc':
            return self._convert_doc(input_path)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
            return [Image.open(input_path)]
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    def _convert_pdf(self, pdf_path: str) -> List[Image.Image]:
        """
        将PDF转换为图片列表
        使用pdf2image进行转换（DPI=200以保证质量和性能平衡）
        转换完成后自动裁剪底部（去除logo），裁剪高度为图片高度的7.8%
        
        Args:
            pdf_path: PDF文件路径
            
        Returns:
            图片对象列表（已裁剪底部）
        """
        if not PDF_SUPPORT:
            raise ImportError("需要安装 pdf2image 库来处理PDF文件")
        
        try:
            # 尝试自动检测 poppler 路径
            poppler_path = self._detect_poppler_path()
            
            # 提高DPI以提升OCR识别质量和与原始文档的一致性
            # 使用200 DPI以获得更好的质量（与原文档样式更接近）
            if poppler_path:
                images = convert_from_path(pdf_path, dpi=200, poppler_path=poppler_path)
            else:
                images = convert_from_path(pdf_path, dpi=200)
            
            # 对每张图片进行底部裁剪（去除logo）
            cropped_images = []
            for img in images:
                cropped_img = self._crop_bottom_logo(img)
                cropped_images.append(cropped_img)
            
            if len(cropped_images) > 0:
                print(f"  [底部裁剪] 已裁剪 {len(cropped_images)} 页的底部logo")
            
            return cropped_images
        except Exception as e:
            raise Exception(f"PDF转换失败: {e}")
    
    def _crop_bottom_logo(self, image: Image.Image) -> Image.Image:
        """
        裁剪图片底部（去除logo）
        裁剪高度为图片高度的5%（0.05）
        
        Args:
            image: PIL图片对象
            
        Returns:
            裁剪后的图片
        """
        if image is None:
            return image
        
        width, height = image.size
        
        # 计算底部裁剪高度
        bottom_crop_height = int(height * 0.05)
        
        # 确保裁剪高度有效
        if bottom_crop_height <= 0 or bottom_crop_height >= height:
            # 裁剪高度无效，返回原图
            return image
        
        # 裁剪底部：从 (0, 0) 到 (width, height - bottom_crop_height)
        new_height = height - bottom_crop_height
        cropped = image.crop((0, 0, width, new_height))
        
        return cropped
    
    def _detect_poppler_path(self) -> str:
        """
        检测 poppler 安装路径
        
        Returns:
            poppler bin 目录路径，如果未找到返回 None
        """
        import platform
        
        if platform.system() == "Windows":
            # 常见安装路径
            possible_paths = [
                r"C:\poppler-25.07.0\Library\bin",
                r"C:\poppler\Library\bin",
                r"C:\Program Files\poppler\bin",
                r"C:\Program Files (x86)\poppler\bin",
            ]
            
            for path in possible_paths:
                pdfinfo_exe = os.path.join(path, "pdfinfo.exe")
                if os.path.exists(pdfinfo_exe):
                    print(f"[自动检测] 找到 poppler: {path}")
                    return path
        
        return None
    
    def extract_docx_text(self, docx_path: str) -> List[str]:
        """
        从DOCX提取纯文本内容
        
        Args:
            docx_path: DOCX文件路径
            
        Returns:
            段落文本列表
        """
        if not DOCX_SUPPORT:
            raise ImportError("需要安装 python-docx 库来处理DOCX文件")
        
        doc = DocxDocument(docx_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return paragraphs
    
    def _convert_docx(self, docx_path: str) -> List[Image.Image]:
        """
        将DOCX转换为图片列表
        使用临时PDF转换方法
        
        Args:
            docx_path: DOCX文件路径
            
        Returns:
            图片对象列表
        """
        if not DOCX_SUPPORT:
            raise ImportError("需要安装 python-docx 库来处理DOCX文件")
        
        # 步骤1: 从DOCX提取纯文本并识别关键词
        try:
            paragraphs = self.extract_docx_text(docx_path)
            print(f"[文本提取] 从DOCX提取了 {len(paragraphs)} 个段落")
            # 这里不处理关键词，留给layout_analyzer处理
        except Exception as e:
            print(f"[警告] 文本提取失败: {e}")
            paragraphs = []
        
        # 优先使用 WPS Office + pdf2image（质量更好，与原文档样式更接近）
        try:
            images = self._convert_with_wps(docx_path)
            if images:
                # 检查是否需要将双页并排拆分成单页
                processed_images = self._split_double_page_layout(images)
                print(f"[WPS] 转换成功，共 {len(processed_images)} 页")
                return processed_images
        except Exception as e:
            print(f"[警告] WPS转换失败: {e}")
            print(f"[降级] 将尝试使用LibreOffice作为备用方案...")
            # 不直接返回，继续尝试LibreOffice
        
        # 备用方案：使用LibreOffice/OpenOffice转换
        try:
            images = self._convert_with_libreoffice(docx_path)
            if images:
                # 检查是否需要将双页并排拆分成单页
                processed_images = self._split_double_page_layout(images)
                print(f"[LibreOffice] 转换成功，共 {len(processed_images)} 页")
                return processed_images
        except Exception as e:
            print(f"[警告] LibreOffice转换失败: {e}")
            pass
        
        # 方法3: 直接提取文本渲染（备用方法）
        # 从DOCX提取文本并渲染为图片
        print("[警告] 使用文本渲染方法，图片和特殊字符将丢失")
        images = self._docx_to_images_render(docx_path)
        # 检查是否需要将双页并排拆分成单页
        processed_images = self._split_double_page_layout(images)
        return processed_images
    
    def _split_double_page_layout(self, images: List[Image.Image]) -> List[Image.Image]:
        """
        将双页并排布局拆分成单页
        通过图片尺寸判断是否需要拆分
        
        Args:
            images: 原始图片列表
            
        Returns:
            处理后的图片列表
        """
        if not images:
            return images
        
        processed_images = []
        
        for idx, image in enumerate(images, 1):
            width, height = image.size
            
            # 判断是否为双页并排布局
            # 标准：宽度 > 高度 且 高度 > 1000像素
            is_double_page = width > height and height > 2000
            
            if is_double_page:
                print(f"[排版] 检测到第{idx}页为双页并排布局 ({width}x{height})，开始拆分...")
                
                # 垂直切分为两张图片
                mid_x = width // 2
                left_image = image.crop((0, 0, mid_x, height))
                right_image = image.crop((mid_x, 0, width, height))
                
                processed_images.append(left_image)
                processed_images.append(right_image)
                print(f"[排版] 已拆分为两张单页")
            else:
                processed_images.append(image)
        
        if len(processed_images) > len(images):
            print(f"[排版] 共拆分 {len(images)} 页为 {len(processed_images)} 页")
        
        return processed_images
    
    def _convert_doc(self, doc_path: str) -> List[Image.Image]:
        """
        将DOC转换为图片列表
        使用COM对象（仅Windows）或转换为DOCX后处理
        
        Args:
            doc_path: DOC文件路径
            
        Returns:
            图片对象列表
        """
        if not DOC_SUPPORT:
            raise ImportError("DOC格式处理需要Windows系统且安装pywin32库")
        
        try:
            # 使用COM对象转换为PDF
            return self._convert_doc_with_com(doc_path)
        except Exception as e:
            # 备用方案：提示用户先转换为DOCX
            raise Exception(f"DOC转换失败: {e}\n提示: 可以将DOC文件转换为DOCX或PDF后处理")
    
    def _convert_with_wps(self, docx_path: str) -> List[Image.Image]:
        """
        使用WPS Office将DOCX文档转换为PDF，再转为图片（修复路径问题）
        
        Args:
            docx_path: DOCX文件路径
            
        Returns:
            图片对象列表
        """
        import tempfile
        import time
        
        wps = None
        doc = None
        
        try:
            # 处理文件路径（确保使用绝对路径）
            if not os.path.isabs(docx_path):
                docx_path = os.path.abspath(docx_path)
            
            # 确保文件存在
            if not os.path.exists(docx_path):
                # 列出当前目录的文件来帮助调试
                current_dir = os.path.dirname(docx_path)
                if os.path.exists(current_dir):
                    files_in_dir = os.listdir(current_dir)
                    docx_files = [f for f in files_in_dir if f.endswith('.docx')]
                    if docx_files:
                        print(f"[WPS] 当前目录下的 .docx 文件: {docx_files}")
                raise FileNotFoundError(f"文件不存在: {docx_path}")
            
            print(f"[WPS] 查找文件: {docx_path}")
            
            # 初始化 WPS Application（使用动态调用，不依赖特定方法）
            try:
                wps = win32com.client.Dispatch("Kwps.Application")  # WPS 文字
                print("[WPS] 使用 Kwps.Application")
            except:
                try:
                    wps = win32com.client.Dispatch("Wps.Application")  # 另一种可能的 WPS 程序名
                    print("[WPS] 使用 Wps.Application")
                except:
                    try:
                        wps = win32com.client.Dispatch("ket.Application")  # WPS 英文版
                        print("[WPS] 使用 ket.Application")
                    except:
                        raise Exception("无法连接到 WPS Office，请确保已安装 WPS")
            
            # 尝试使用动态调用（COM对象的方法可能不通过hasattr检测到）
            # 直接尝试调用，捕获异常
            
            wps.Visible = False
            wps.DisplayAlerts = False  # 禁用警告提示
            
            # 使用临时目录（不使用上下文管理器，避免提前删除）
            import uuid
            temp_dir = os.path.join(tempfile.gettempdir(), f"wps_convert_{uuid.uuid4().hex}")
            os.makedirs(temp_dir, exist_ok=True)
            
            try:
                # 转换 DOCX 为 PDF
                pdf_path = os.path.join(temp_dir, "temp.pdf")
                
                print("[WPS] 正在打开文档...")
                print(f"[WPS] 文档路径: {docx_path}")
                
                # 打开文档（确保路径格式正确）
                # 将路径转换为Windows格式（如果包含空格等特殊字符需要特殊处理）
                docx_path_normalized = docx_path.replace('/', '\\')
                
                try:
                    doc = wps.Documents.Open(docx_path_normalized)
                except Exception as e:
                    # 如果失败，尝试原始路径
                    print(f"[WPS] 尝试原始路径失败: {e}，尝试其他方式")
                    doc = wps.Documents.Open(docx_path)
                
                print("[WPS] 正在转换为 PDF...")
                
                # 先列出文档对象的所有可用方法（用于调试）
                if self.config and self.config.mode == 'debug':
                    try:
                        methods = [method for method in dir(doc) if not method.startswith('_')]
                        print(f"[WPS调试] 文档对象可用方法（前20个）: {methods[:20]}")
                    except:
                        pass
                
                # 导出为 PDF（使用不同的方法尝试）
                # 注意：COM对象的方法可能不通过hasattr检测到，直接尝试调用
                export_success = False
                
                # 方法1：尝试使用 ExportAsFixedFormat（Microsoft Word标准方法）
                if not export_success:
                    try:
                        doc.ExportAsFixedFormat(
                            OutputFileName=pdf_path,
                            ExportFormat=17,  # 17 表示 PDF 格式
                            OpenAfterExport=False
                        )
                        export_success = True
                        print("[WPS] 使用 ExportAsFixedFormat 成功")
                    except AttributeError:
                        print("[WPS] ExportAsFixedFormat 方法不存在")
                    except Exception as e1:
                        # 尝试不同的参数组合
                        try:
                            doc.ExportAsFixedFormat(pdf_path, 17)
                            export_success = True
                            print("[WPS] 使用简化参数成功")
                        except Exception as e1b:
                            print(f"[WPS] ExportAsFixedFormat 失败: {e1}, {e1b}")
                
                # 方法1b：尝试使用 ExportAsFixedFormat2（某些版本的替代方法）
                if not export_success:
                    try:
                        doc.ExportAsFixedFormat2(
                            OutputFileName=pdf_path,
                            ExportFormat=17,
                            OpenAfterExport=False
                        )
                        export_success = True
                        print("[WPS] 使用 ExportAsFixedFormat2 成功")
                    except (AttributeError, Exception) as e:
                        if "AttributeError" not in str(type(e)):
                            print(f"[WPS] ExportAsFixedFormat2 失败: {e}")
                
                # 方法2：尝试使用WPS特有的Export方法
                if not export_success:
                    try:
                        doc.Export(pdf_path, "pdf")  # 可能使用字符串格式
                        time.sleep(0.5)
                        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                            export_success = True
                            print("[WPS] 使用 Export(字符串) 成功")
                    except AttributeError:
                        try:
                            doc.Export(pdf_path, 17)  # 或使用数字格式
                            time.sleep(0.5)
                            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                                export_success = True
                                print("[WPS] 使用 Export(数字) 成功")
                        except (AttributeError, Exception) as e2:
                            if "AttributeError" not in str(type(e2)):
                                print(f"[WPS] Export 失败: {e2}")
                    except Exception as e:
                        print(f"[WPS] Export 尝试失败: {e}")
                
                # 方法3：尝试使用 SaveAs（不同参数组合）
                if not export_success:
                    try:
                        # 尝试多种文件格式代码和参数传递方式
                        format_codes = [
                            (17, "FileFormat"),  # Microsoft Word PDF格式
                            (32, "FileFormat"),  # 可能的WPS PDF格式
                            (0, "FileFormat"),   # 默认格式
                            (17, None),         # 位置参数
                        ]
                        
                        for fmt_code, param_name in format_codes:
                            try:
                                if param_name:
                                    doc.SaveAs(pdf_path, **{param_name: fmt_code})
                                else:
                                    doc.SaveAs(pdf_path, fmt_code)
                                
                                # 等待并检查文件
                                time.sleep(0.5)
                                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                                    export_success = True
                                    print(f"[WPS] 使用 SaveAs (格式码={fmt_code}) 成功")
                                    break
                            except AttributeError:
                                break  # 如果SaveAs不存在，跳过所有尝试
                            except Exception as e:
                                continue
                        if not export_success:
                            print("[WPS] SaveAs 方法不存在或无法使用")
                    except AttributeError:
                        print("[WPS] SaveAs 方法不存在")
                    except Exception as e2:
                        print(f"[WPS] SaveAs 方法失败: {e2}")
                
                # 方法4：尝试通过Application级别的Export方法
                if not export_success:
                    try:
                        wps.ExportDocument(docx_path_normalized, pdf_path, "pdf")
                        time.sleep(1)
                        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                            export_success = True
                            print("[WPS] 使用 Application.ExportDocument 成功")
                    except AttributeError:
                        try:
                            wps.Export(docx_path_normalized, pdf_path, "pdf")
                            time.sleep(1)
                            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                                export_success = True
                                print("[WPS] 使用 Application.Export 成功")
                        except (AttributeError, Exception) as e4:
                            if "AttributeError" not in str(type(e4)):
                                print(f"[WPS] Application级别导出失败: {e4}")
                    except Exception as e4:
                        print(f"[WPS] Application级别导出失败: {e4}")
                
                # 方法5：尝试使用ActiveDocument.SaveAs2（WPS可能支持）
                if not export_success:
                    try:
                        wps.ActiveDocument.SaveAs2(pdf_path, FileFormat=17)
                        time.sleep(1)
                        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                            export_success = True
                            print("[WPS] 使用 ActiveDocument.SaveAs2 成功")
                    except (AttributeError, Exception) as e5:
                        if "AttributeError" not in str(type(e5)):
                            print(f"[WPS] SaveAs2 失败: {e5}")
                
                # 如果所有方法都失败
                if not export_success:
                    # 列出文档对象的所有方法供调试
                    print("[WPS错误] 所有PDF导出方法都失败")
                    print("[WPS调试] 建议：")
                    print("  1. 检查WPS Office版本是否支持COM接口")
                    print("  2. 尝试手动在WPS中打开文件并另存为PDF")
                    print("  3. 使用LibreOffice作为替代方案")
                    raise Exception("所有PDF导出方法都失败。WPS COM接口可能不支持自动导出PDF")
                
                print("[WPS] 正在关闭文档...")
                # 关闭文档
                doc.Close(SaveChanges=False)
                doc = None
                
                # 等待文件写入完成
                time.sleep(2)  # 增加等待时间
                
                # 检查 PDF 文件是否生成
                if not os.path.exists(pdf_path):
                    raise Exception(f"PDF 文件生成失败，路径: {pdf_path}")
                
                print(f"[WPS] PDF 已生成: {pdf_path}")
                print("[WPS] 正在将 PDF 转换为图片...")
                
                # 将 PDF 转换为图片（使用更高的DPI以获得更好的质量）
                images = self._convert_pdf(pdf_path)
                
                # 保存转换后的图片到output_test目录以供检查
                output_test_dir = "output_test"
                os.makedirs(output_test_dir, exist_ok=True)
                for idx, img in enumerate(images, 1):
                    base_name = Path(docx_path).stem
                    output_path = os.path.join(output_test_dir, f"{base_name}_page{idx}.png")
                    img.save(output_path, 'PNG', quality=95)
                    print(f"[WPS] 保存测试图片: {output_path}")
                
                return images
                
            finally:
                # 清理临时目录
                try:
                    import shutil
                    shutil.rmtree(temp_dir)
                except:
                    pass
                
        except Exception as e:
            error_msg = str(e)
            # 提供更详细的错误信息
            if "ExportAsFixedFormat" in error_msg or "Open." in error_msg:
                print(f"[WPS] 错误详情: {error_msg}")
                print(f"[WPS] 提示: 这可能是WPS版本或COM对象接口的问题")
                print(f"[WPS] 建议: 检查WPS Office是否已正确安装，或尝试使用LibreOffice")
            raise Exception(f"WPS转换失败: {error_msg}")
        finally:
            # 优雅地清理资源
            try:
                if doc is not None:
                    try:
                        doc.Close(SaveChanges=False)
                    except:
                        pass
            except:
                pass
            try:
                if wps is not None:
                    try:
                        wps.Quit()
                    except:
                        pass
            except:
                pass
    
    def _convert_with_libreoffice(self, file_path: str) -> List[Image.Image]:
        """
        使用LibreOffice将Office文档转换为PDF，再转为图片
        
        Args:
            file_path: 文件路径
            
        Returns:
            图片对象列表
        """
        import tempfile
        import subprocess
        
        # 查找LibreOffice可执行文件
        import shutil
        
        # 优先检查系统PATH中是否已安装
        libreoffice_path = shutil.which('soffice')
        
        if not libreoffice_path:
            # 尝试常见安装路径
            common_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
            for path in common_paths:
                if os.path.exists(path):
                    libreoffice_path = path
                    print(f"[LibreOffice] 找到安装路径: {libreoffice_path}")
                    break
        
        if not libreoffice_path:
            raise FileNotFoundError("未找到LibreOffice，请安装LibreOffice以支持Office文档")
        
        # 创建临时目录（不使用上下文管理器，避免提前删除）
        import uuid
        temp_dir = os.path.join(tempfile.gettempdir(), f"libreoffice_{uuid.uuid4().hex}")
        os.makedirs(temp_dir, exist_ok=True)
        
        try:
            print(f"[LibreOffice] 开始转换: {file_path}")
            base_name = Path(file_path).stem
            
            # 方法1: 转换为PDF再转为图片（推荐，支持多页且格式保持好）
            result = subprocess.run([
                libreoffice_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', temp_dir,
                file_path
            ], check=True, capture_output=True, text=True, timeout=120)
            
            # 获取生成的PDF路径
            pdf_file = os.path.join(temp_dir, base_name + '.pdf')
            
            if os.path.exists(pdf_file):
                print(f"[LibreOffice] PDF转换成功: {pdf_file}")
                # 使用现有的PDF转换功能
                images = self._convert_pdf(pdf_file)
                print(f"[LibreOffice] 转换成功，共 {len(images)} 页")
                
                # 保存转换后的图片到output_test目录以供检查
                output_test_dir = "output_test"
                os.makedirs(output_test_dir, exist_ok=True)
                for idx, img in enumerate(images, 1):
                    output_path = os.path.join(output_test_dir, f"{base_name}_page{idx}.png")
                    img.save(output_path)
                    print(f"[LibreOffice] 保存测试图片: {output_path}")
                
                return images
            
            print(f"[LibreOffice] 警告: 未生成PDF文件")
            print(f"[LibreOffice] STDOUT: {result.stdout}")
            print(f"[LibreOffice] STDERR: {result.stderr}")
            raise FileNotFoundError("LibreOffice PDF转换失败")
        finally:
            # 清理临时目录
            import shutil
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
    
    def _docx_to_images_render(self, docx_path: str) -> List[Image.Image]:
        """
        从DOCX提取文本并渲染为图片
        这是一个简单的备用方法
        
        Args:
            docx_path: DOCX文件路径
            
        Returns:
            图片对象列表
        """
        doc = DocxDocument(docx_path)
        
        # 提取所有段落文本
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        if not paragraphs:
            raise ValueError("DOCX文件中未找到文本内容")
        
        # 简单的文本渲染为图片
        from PIL import ImageDraw, ImageFont
        
        # 创建图片
        img_width = 800
        line_height = 40
        margin = 50
        img_height = len(paragraphs) * line_height + margin * 2
        
        # 尝试加载字体
        try:
            if os.name == 'nt':  # Windows
                font = ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", 20)
            else:
                font = ImageFont.load_default()
        except:
            font = ImageFont.load_default()
        
        image = Image.new('RGB', (img_width, img_height), color='white')
        draw = ImageDraw.Draw(image)
        
        y = margin
        for para in paragraphs:
            # 简单的文本换行
            words = para
            if len(words) > 40:
                words = words[:37] + "..."
            draw.text((margin, y), words, fill='black', font=font)
            y += line_height
        
        return [image]
    
    def _convert_doc_with_com(self, doc_path: str) -> List[Image.Image]:
        """
        使用COM对象将DOC转换为PDF（仅Windows）
        
        Args:
            doc_path: DOC文件路径
            
        Returns:
            图片对象列表
        """
        import tempfile
        
        # 使用Word COM对象转换为PDF
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                pdf_file = os.path.join(temp_dir, Path(doc_path).stem + '.pdf')
                
                # 打开并转换
                doc = word.Documents.Open(doc_path)
                doc.SaveAs(pdf_file, FileFormat=17)  # 17 = PDF格式
                doc.Close()
                
                # 使用现有的PDF转换功能
                return self._convert_pdf(pdf_file)
        finally:
            word.Quit()
    
    def get_file_info(self, input_path: str) -> dict:
        """
        获取文件基本信息
        
        Args:
            input_path: 输入文件路径
            
        Returns:
            文件信息字典
        """
        file_path = Path(input_path)
        
        info = {
            "path": str(input_path),
            "name": file_path.name,
            "stem": file_path.stem,
            "extension": file_path.suffix,
            "size": file_path.stat().st_size if file_path.exists() else 0
        }
        
        return info

